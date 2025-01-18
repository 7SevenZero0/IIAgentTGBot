from aiogram.types import Message, ContentType, FSInputFile
from aiogram import Router, F
from app.database.models import async_session, AudioRecording
from aiogram.filters import CommandStart, Command
from sqlalchemy.sql import text
from datetime import datetime
from app.keyboards import generate_audio_keyboard, main_menu_keyboard
from pydub import AudioSegment
import os
import speech_recognition as sr
from docx import Document

router = Router()

if not os.path.exists('uploads'):
    os.makedirs('uploads')

def convert_mp3_to_wav(mp3_file, wav_file):
    print(f"Конвертируем файл {mp3_file} в формат wav...")
    audio = AudioSegment.from_mp3(mp3_file)
    audio.export(wav_file, format="wav")
    print(f"Файл {mp3_file} успешно конвертирован в {wav_file}")

def recognize_audio(wav_file):
    print(f"Распознаем речь в файле {wav_file}...")
    recognizer = sr.Recognizer()
    
    with sr.AudioFile(wav_file) as source:
        audio = recognizer.record(source)
    
    try:
        text = recognizer.recognize_google(audio, language='ru-RU')
        print("Речь успешно распознана")
        return text
    except sr.UnknownValueError:
        print(f"Не удалось распознать речь в файле {wav_file}")
        return "Не удалось распознать речь"
    except sr.RequestError as e:
        print(f"Ошибка запроса к сервису Google Speech Recognition: {e}")
        return f"Ошибка запроса к сервису Google Speech Recognition: {e}"

@router.message(CommandStart())
async def cmd_start(message: Message):
    print(f"Пользователь {message.from_user.id} нажал /start")
    await message.answer('Добро пожаловать! Выберите действие из меню ниже:', reply_markup=main_menu_keyboard())

@router.message(Command('help'))
async def cmd_help(message: Message):
    print(f"Пользователь {message.from_user.id} запросил помощь")
    await message.answer('Доступные команды:\n- Загрузить аудиозапись\n- Получить транскрипцию\n- Сформировать протокол\n- Просмотреть загруженные аудиофайлы')

@router.message(F.text == 'Загрузить аудиозапись')
async def upload_audio(message: Message):
    await message.answer('Отправьте аудиозапись, чтобы начать обработку.')

@router.message(F.content_type == ContentType.AUDIO)
async def handle_audio(message: Message):
    print(f"Пользователь {message.from_user.id} загрузил аудиофайл")
    audio = message.audio
    file_id = audio.file_id
    file = await message.bot.get_file(file_id)

    file_name = audio.file_name if audio.file_name else f"{file_id}.mp3"
    
    file_name = file_name.replace(" ", "_").replace("/", "_")
    file_path = os.path.join("uploads", file_name)

    await message.bot.download_file(file.file_path, file_path)

    async with async_session() as session:
        audio_recording = AudioRecording(
            record_path=file_path,
            date_time=datetime.utcnow(),
            user_id=message.from_user.id
        )
        session.add(audio_recording)
        await session.commit()

    print(f"Аудиофайл {file_name} успешно загружен и сохранен в базу данных")
    await message.answer(f'Аудиозапись "{file_name}" успешно загружена и сохранена.')

@router.message(F.text == 'Получить транскрипцию')
async def get_transcription(message: Message):
    print(f"Пользователь {message.from_user.id} запросил транскрипцию")
    async with async_session() as session:
        user_id = message.from_user.id
        result = await session.execute(
            text("SELECT record_path FROM audio_recordings WHERE user_id = :user_id ORDER BY date_time DESC LIMIT 1"),
            {"user_id": user_id}
        )
        audio_files = result.scalars().all()

    if audio_files:
        response = "Выберите аудиофайл для транскрипции:"
        keyboard = generate_audio_keyboard(audio_files)
        print(f"Пользователю {message.from_user.id} отправлены аудиофайлы для транскрипции")
        await message.answer(response, reply_markup=keyboard)
    else:
        print(f"Пользователь {message.from_user.id} не имеет загруженных аудиофайлов")
        await message.answer("У вас нет загруженных аудиофайлов для транскрипции.")

@router.message(F.text == 'Просмотреть мои записи')
async def view_user_audio(message: Message):
    async with async_session() as session:
        user_id = message.from_user.id
        result = await session.execute(
            text("SELECT record_path FROM audio_recordings WHERE user_id = :user_id"),
            {"user_id": user_id}
        )
        audio_files = result.scalars().all()

    if audio_files:
        response = "Ваши аудиозаписи:\n"
        keyboard = generate_audio_keyboard(audio_files)
        await message.answer(response, reply_markup=keyboard)
    else:
        response = "У вас пока нет загруженных аудиозаписей."
        await message.answer(response)

@router.message(F.text == 'Вернуться')
async def back(message: Message):
    print(f"Пользователь {message.from_user.id} вернулся в главное меню")
    await message.answer('Вы вернулись на главную', reply_markup=main_menu_keyboard())

@router.message(F.text)
async def selected_audio_for_transcription(message: Message):
    audio_file_name = message.text
    print(f"Пользователь {message.from_user.id} выбрал аудиофайл {audio_file_name} для транскрипции")
    async with async_session() as session:
        result = await session.execute(
            text("SELECT record_path FROM audio_recordings WHERE record_path LIKE :file_name LIMIT 1"),
            {"file_name": f"%{audio_file_name}%"}
        )
        audio_file = result.scalar()

    if audio_file:
        wav_file_path = audio_file.replace('.mp3', '.wav')
        convert_mp3_to_wav(audio_file, wav_file_path)

        transcription = recognize_audio(wav_file_path)

        print(f"Транскрипция для аудиофайла {audio_file_name}: {transcription}")
        await message.answer(f"Транскрипция для {audio_file_name}:\n{transcription}", reply_markup=main_menu_keyboard())

        report_file_path = generate_docx_report(audio_file_name, transcription, message.from_user.id)
        file = FSInputFile(report_file_path)
        await message.answer_document(file, caption=f"Отчет по аудиофайлу {audio_file_name}")
    else:
        print(f"Аудиофайл {audio_file_name} не найден для пользователя {message.from_user.id}")
        await message.answer(f"Не удалось найти аудиофайл с именем {audio_file_name}. Попробуйте снова.", reply_markup=main_menu_keyboard())

def generate_docx_report(audio_file_name: str, transcription: str, user_id: int) -> str:

    file_path = f"{audio_file_name}_report_{user_id}.docx"
    doc = Document()
    doc.add_heading(f"Отчет по транскрипции аудиофайла {audio_file_name}", level=1)
    doc.add_paragraph(f"Пользователь ID: {user_id}")
    doc.add_paragraph(f"Название аудиофайла: {audio_file_name}")
    doc.add_heading("Транскрипция:", level=2)
    doc.add_paragraph(transcription)
    doc.save(file_path)
    return file_path